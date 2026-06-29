import os
import pandas as pd
from docx import Document

INPUT_EXCEL = "참석자명단.xlsx"
TEMPLATE_DOCX = "공문양식_템플릿.docx"
OUTPUT_DIR = "생성결과"

os.makedirs(OUTPUT_DIR, exist_ok=True)

df = pd.read_excel(INPUT_EXCEL, sheet_name="참석자명단")

def replace_text(doc, mapping):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in mapping.items():
                run.text = run.text.replace("{{" + key + "}}", str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in mapping.items():
                            run.text = run.text.replace("{{" + key + "}}", str(value))

for _, row in df.iterrows():
    doc = Document(TEMPLATE_DOCX)
    mapping = {
        "이름": row["이름"],
        "부서": row["부서"],
        "직급": row["직급"],
        "교육일자": pd.to_datetime(row["교육일자"]).strftime("%Y-%m-%d"),
        "과정명": row["과정명"],
        "교육장소": row["교육장소"],
    }
    replace_text(doc, mapping)
    filename = f'{mapping["이름"]}_참석요청.docx'
    doc.save(os.path.join(OUTPUT_DIR, filename))

print(f"완료: {len(df)}개 워드 파일 생성")
